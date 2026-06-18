"""
ARIA — Document Processor

Read and write document formats:
- PDF (read only, using PyMuPDF)
- DOCX (read/write, using python-docx)
- XLSX (read/write, using openpyxl)

Inspired by the Anthropic skills document processing patterns.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional


SUPPORTED_EXTENSIONS = {
    ".pdf": "PDF Document",
    ".docx": "Word Document",
    ".xlsx": "Excel Spreadsheet",
    ".txt": "Text File",
    ".md": "Markdown File",
    ".py": "Python File",
    ".json": "JSON File",
    ".yaml": "YAML File",
    ".yml": "YAML File",
}


def read_document(filepath: str) -> Dict[str, Any]:
    """
    Read a document and return its content.

    Args:
        filepath: Path to the file

    Returns:
        Dict with 'content' (str), 'type' (str), 'pages' (int for PDF),
        'sheets' (list for XLSX), or 'error' if failed
    """
    path = Path(filepath)
    if not path.exists():
        return {"error": f"File not found: {filepath}"}

    ext = path.suffix.lower()

    if ext == ".pdf":
        return _read_pdf(path)
    elif ext == ".docx":
        return _read_docx(path)
    elif ext == ".xlsx":
        return _read_xlsx(path)
    elif ext in (".txt", ".md", ".py", ".json", ".yaml", ".yml"):
        return _read_text(path)
    else:
        return {"error": f"Unsupported format: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"}


def write_document(filepath: str, content: Any, format: Optional[str] = None) -> Dict[str, Any]:
    """
    Write content to a document file.

    Args:
        filepath: Path to save the file
        content: Content to write (str for PDF/DOCX, list of dicts for XLSX)
        format: Override format (pdf, docx, xlsx). If None, inferred from extension.

    Returns:
        Dict with 'path' (str), 'type' (str), or 'error' if failed
    """
    path = Path(filepath)
    ext = format or path.suffix.lower()

    if ext == ".docx" or format == "docx":
        return _write_docx(path, content)
    elif ext == ".xlsx" or format == "xlsx":
        return _write_xlsx(path, content)
    else:
        return {"error": f"Write not supported for {ext}. Supported: .docx, .xlsx"}


# ─── PDF Reader ────────────────────────────────────────────────────────────

def _read_pdf(path: Path) -> Dict[str, Any]:
    """Read a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {"error": "PyMuPDF not installed. Run: pip install PyMuPDF"}

    try:
        doc = fitz.open(path)
        page_count = len(doc)
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages.append({"page": i + 1, "text": text.strip()})

        full_text = "\n\n".join(p["text"] for p in pages)
        doc.close()

        return {
            "content": full_text,
            "type": "pdf",
            "pages": page_count,
            "page_count": len(pages),
        }
    except Exception as e:
        return {"error": f"Failed to read PDF: {e}"}


# ─── DOCX Reader/Writer ───────────────────────────────────────────────────

def _read_docx(path: Path) -> Dict[str, Any]:
    """Read a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    try:
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        # Extract tables too
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(rows)

        return {
            "content": full_text,
            "type": "docx",
            "paragraphs": len(paragraphs),
            "tables": tables,
        }
    except Exception as e:
        return {"error": f"Failed to read DOCX: {e}"}


def _write_docx(path: Path, content: str) -> Dict[str, Any]:
    """Write content to a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    try:
        doc = Document()
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[2:], level=3)
            else:
                doc.add_paragraph(line)

        doc.save(path)
        return {"path": str(path), "type": "docx"}
    except Exception as e:
        return {"error": f"Failed to write DOCX: {e}"}


# ─── Text Reader (for .txt, .md, .py, .json, .yaml) ──────────────────────

def _read_text(path: Path) -> Dict[str, Any]:
    """Read a plain text file."""
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        return {
            "content": content.strip(),
            "type": path.suffix[1:] if path.suffix else "text",
            "lines": content.count("\n") + 1,
        }
    except Exception as e:
        return {"error": f"Failed to read file: {e}"}


# ─── XLSX Reader/Writer ───────────────────────────────────────────────────

def _read_xlsx(path: Path) -> Dict[str, Any]:
    """Read an XLSX file using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(c) if c is not None else "" for c in row])
            sheets.append({"name": sheet_name, "rows": rows})

        wb.close()

        return {
            "content": "\n".join(
                f"Sheet: {s['name']}\n" + "\n".join(" | ".join(r) for r in s["rows"][:50])
                for s in sheets
            ),
            "type": "xlsx",
            "sheets": [s["name"] for s in sheets],
        }
    except Exception as e:
        return {"error": f"Failed to read XLSX: {e}"}


def _write_xlsx(path: Path, data: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Write data to an XLSX file.

    Args:
        path: Output path
        data: List of dicts, each dict is a row. Keys become headers.
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"

        if data:
            # Headers
            headers = list(data[0].keys())
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")

            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.fill = header_fill
                cell.font = header_font

            # Data rows
            for row_idx, row_data in enumerate(data, 2):
                for col_idx, header in enumerate(headers, 1):
                    ws.cell(row=row_idx, column=col_idx, value=row_data.get(header, ""))

            # Auto-adjust column widths
            for col in range(1, len(headers) + 1):
                max_len = len(str(headers[col - 1]))
                for row in range(2, len(data) + 2):
                    cell_val = ws.cell(row=row, column=col).value
                    if cell_val:
                        max_len = max(max_len, len(str(cell_val)))
                ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = min(max_len + 2, 50)

        wb.save(path)
        wb.close()
        return {"path": str(path), "type": "xlsx"}
    except Exception as e:
        return {"error": f"Failed to write XLSX: {e}"}
